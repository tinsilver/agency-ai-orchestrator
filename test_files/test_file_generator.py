"""
Generate test files for file upload feature testing.
Creates DOCX, PDF, and image files with sample content.
"""
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw, ImageFont

def create_test_docx():
    """Create a test Word document with a sample change request."""
    doc = Document()
    doc.add_heading('Website Change Request', 0)
    
    doc.add_paragraph('Client: The Business Beanstalk')
    doc.add_paragraph('Date: 2026-02-01')
    doc.add_paragraph()
    
    doc.add_heading('Request Details', 1)
    doc.add_paragraph(
        'Please add a contact form to the About page with the following fields:'
    )
    
    # Add a bulleted list
    doc.add_paragraph('Name (required)', style='List Bullet')
    doc.add_paragraph('Email (required)', style='List Bullet')
    doc.add_paragraph('Phone Number (optional)', style='List Bullet')
    doc.add_paragraph('Message (required, text area)', style='List Bullet')
    doc.add_paragraph('Submit button with primary brand color', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('Additional Notes', 1)
    doc.add_paragraph(
        'The form should be positioned below the team member section. '
        'Include form validation and display a success message after submission.'
    )
    
    filepath = os.path.join('test_files', 'test.docx')
    doc.save(filepath)
    print(f"✅ Created {filepath}")
    return filepath

def create_test_pdf():
    """Create a test PDF with technical specifications."""
    filepath = os.path.join('test_files', 'test.pdf')
    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, "Technical Specification Document")
    
    # Content
    c.setFont("Helvetica", 12)
    y = height - 100
    
    lines = [
        "Project: Homepage CTA Button Addition",
        "",
        "Requirements:",
        "1. Button text: 'View Our Services'",
        "2. Position: Directly below hero section (margin: 40px)",
        "3. Styling:",
        "   - Background: Primary Red (#ff0000)",
        "   - Text: White (#ffffff)",
        "   - Border radius: 8px",
        "   - Padding: 16px 32px",
        "   - Font size: 18px",
        "4. Hover effect: Background changes to Secondary Blue (#0000ff)",
        "5. Link destination: /services page",
        "",
        "Success Criteria:",
        "- Button is visible on all viewport sizes",
        "- Click triggers navigation to Services page",
        "- Meets WCAG AA contrast requirements"
    ]
    
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
        if y < 50:  # Start new page if needed
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 50
    
    c.save()
    print(f"✅ Created {filepath}")
    return filepath

def create_test_image():
    """Create a simple wireframe mockup image."""
    # Create a simple mockup image
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)
    
    # Draw a simple wireframe
    # Header
    draw.rectangle([50, 50, 750, 120], outline='black', width=2)
    draw.text((360, 75), "LOGO", fill='black')
    
    # Hero Section
    draw.rectangle([50, 140, 750, 350], outline='black', width=2)
    draw.text((300, 230), "Hero Section with Main Image", fill='gray')
    
    # CTA Button (highlighted in red)
    draw.rectangle([300, 380, 500, 440], fill='#ff0000', outline='black', width=2)
    draw.text((325, 400), "View Our Services", fill='white')
    
    # Content area
    draw.rectangle([50, 470, 750, 550], outline='black', width=2)
    draw.text((320, 500), "Content Area", fill='gray')
    
    filepath = os.path.join('test_files', 'test.jpg')
    img.save(filepath, 'JPEG')
    print(f"✅ Created {filepath}")
    return filepath

def create_wireframe_png():
    """Create a PNG wireframe alternative."""
    img = Image.new('RGB', (1000, 800), color='#f5f5f5')
    draw = ImageDraw.Draw(img)
    
    # Title
    draw.text((50, 30), "Contact Form Wireframe", fill='black')
    
    # Form outline
    draw.rectangle([100, 80, 900, 700], outline='black', width=3)
    
    # Form fields
    fields = [
        ("Name:", 130),
        ("Email:", 230),
        ("Phone:", 330),
        ("Message:", 430)
    ]
    
    for label, y_pos in fields:
        draw.text((120, y_pos), label, fill='black')
        draw.rectangle([120, y_pos + 25, 880, y_pos + 65], outline='gray', width=2)
    
    # Message field (larger)
    draw.rectangle([120, 455, 880, 595], outline='gray', width=2)
    
    # Submit button
    draw.rectangle([120, 630, 280, 680], fill='#ff0000', outline='black', width=2)
    draw.text((155, 650), "Submit", fill='white')
    
    filepath = os.path.join('test_files', 'mockwireframe.png')
    img.save(filepath, 'PNG')
    print(f"✅ Created {filepath}")
    return filepath

if __name__ == "__main__":
    print("Generating test files...")
    create_test_docx()
    create_test_pdf()
    create_test_image()
    create_wireframe_png()
    print("\n✅ All test files created successfully!")
